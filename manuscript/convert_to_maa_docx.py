#!/usr/bin/env python3
"""
convert_to_maa_docx.py
======================
Converts paper_a_primary_unesco.tex → a MAA-formatted .docx

MAA formatting rules (from https://www.maajournal.com/index.php/maa/authors):
  - Font:         Times New Roman, 12pt body
  - Line spacing: Single-spaced
  - Margins:      ~2.5 cm all sides (A4)
  - Title:        Bold, centered, 14pt
  - Authors:      Centered, 12pt
  - Abstract:     Italic label "Abstract:", ≤250 words, indented block
  - Keywords:     Bold label "Keywords:", max 8, no title words
  - Headings:     H1 = Bold 12pt; H2 = Bold Italic 12pt; H3 = Italic 12pt
  - Body text:    12pt, justified, single-spaced
  - References:   Alphabetical, hanging indent, MAA author-year format
  - Figures/tables: Inline in body

Steps:
  1. Compile the .tex to PDF (resolves all macros)
  2. Use pandoc to convert .tex → .docx with a custom reference.docx
  3. Post-process the .docx with python-docx to enforce MAA styles

Usage:
    python3 convert_to_maa_docx.py

Output:
    manuscript/paper_a_primary_maa.docx
"""

import subprocess
import sys
import os
from pathlib import Path

# ── Paths ──────────────────────────────────────────────────────────────────────
MANUSCRIPT_DIR = Path(__file__).parent
TEX_FILE       = MANUSCRIPT_DIR / "paper_a_primary_unesco.tex"
DOCX_OUT       = MANUSCRIPT_DIR / "paper_a_primary_maa.docx"
REF_DOCX       = MANUSCRIPT_DIR / "maa_reference.docx"
PDF_OUT        = MANUSCRIPT_DIR / "paper_a_primary_unesco.pdf"

# ── Step 1: Compile LaTeX to resolve all macros ───────────────────────────────
def compile_latex():
    print("Step 1: Compiling LaTeX (resolving all macros)...")
    for run in range(2):  # two passes for cross-references
        result = subprocess.run(
            ["pdflatex", "-interaction=nonstopmode",
             "-output-directory", str(MANUSCRIPT_DIR),
             str(TEX_FILE)],
            capture_output=True, text=True, cwd=MANUSCRIPT_DIR
        )
        if result.returncode != 0:
            # Show last 30 lines of log for diagnosis
            log_lines = result.stdout.splitlines()[-30:]
            print("  WARNING: pdflatex returned non-zero. Last log lines:")
            for line in log_lines:
                print("  ", line)
        else:
            print(f"  Pass {run+1} OK.")
    if PDF_OUT.exists():
        print(f"  PDF compiled: {PDF_OUT}")
    else:
        print("  ERROR: PDF not found after compilation. Check .tex for errors.")
        sys.exit(1)

# ── Step 2: Build a MAA-styled reference.docx ─────────────────────────────────
def build_reference_docx():
    """
    Creates a reference.docx that encodes MAA styles.
    pandoc uses this as a style template when generating the output .docx.
    """
    print("Step 2: Building MAA reference.docx style template...")
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy
    except ImportError:
        print("  ERROR: python-docx not installed. Run: pip3 install python-docx")
        sys.exit(1)

    # Start from pandoc's default reference.docx so all styles exist
    result = subprocess.run(
        ["pandoc", "--print-default-data-file", "reference.docx"],
        capture_output=True
    )
    if result.returncode != 0:
        print("  Could not get pandoc default reference.docx, creating from scratch.")
        doc = Document()
    else:
        with open(REF_DOCX, "wb") as f:
            f.write(result.stdout)
        doc = Document(REF_DOCX)

    # ── Page setup: A4, 2.5 cm margins ──────────────────────────────────────
    from docx.shared import Cm
    for section in doc.sections:
        section.page_width  = Cm(21.0)   # A4
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    def set_style(style_name, font_name="Times New Roman", font_size=12,
                  bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=6, line_spacing=None):
        try:
            style = doc.styles[style_name]
        except KeyError:
            return  # style not present in reference doc, skip
        style.font.name       = font_name
        style.font.size       = Pt(font_size)
        style.font.bold       = bold
        style.font.italic     = italic
        pf = style.paragraph_format
        pf.alignment          = alignment
        pf.space_before       = Pt(space_before)
        pf.space_after        = Pt(space_after)
        if line_spacing is not None:
            from docx.shared import Pt as Pts
            from docx.enum.text import WD_LINE_SPACING
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Body text
    set_style("Normal",        font_size=12, bold=False, italic=False,
              space_before=0, space_after=6)
    # Title: 14pt bold centered
    set_style("Title",         font_size=14, bold=True,  italic=False,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    # Heading 1: Bold 12pt (Introduction, Methods, Results, Discussion)
    set_style("Heading 1",     font_size=12, bold=True,  italic=False,
              alignment=WD_ALIGN_PARAGRAPH.LEFT,  space_before=12, space_after=6)
    # Heading 2: Bold Italic 12pt
    set_style("Heading 2",     font_size=12, bold=True,  italic=True,
              alignment=WD_ALIGN_PARAGRAPH.LEFT,  space_before=6,  space_after=3)
    # Heading 3: Italic 12pt
    set_style("Heading 3",     font_size=12, bold=False, italic=True,
              alignment=WD_ALIGN_PARAGRAPH.LEFT,  space_before=6,  space_after=3)
    # Abstract
    set_style("Abstract",      font_size=12, bold=False, italic=False,
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6)
    # Block Text (for abstract indentation)
    set_style("Block Text",    font_size=12, bold=False, italic=False,
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    doc.save(REF_DOCX)
    print(f"  Reference.docx saved: {REF_DOCX}")


# ── Step 3: Convert .tex → .docx via pandoc ───────────────────────────────────
def run_pandoc():
    print("Step 3: Running pandoc (.tex → .docx)...")
    cmd = [
        "pandoc",
        str(TEX_FILE),
        "--from=latex",
        "--to=docx",
        f"--reference-doc={REF_DOCX}",
        "--standalone",
        "--wrap=none",
        "--resource-path=" + str(MANUSCRIPT_DIR),
        "--output=" + str(DOCX_OUT),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, cwd=MANUSCRIPT_DIR)
    if result.returncode != 0:
        print("  pandoc STDERR:", result.stderr[-2000:])
        # non-fatal: pandoc often warns but still produces output
    if DOCX_OUT.exists():
        print(f"  .docx created: {DOCX_OUT}")
    else:
        print("  ERROR: pandoc did not produce output file.")
        sys.exit(1)


# ── Step 4: Post-process with python-docx to enforce MAA styles ───────────────
def post_process_docx():
    print("Step 4: Post-processing .docx for MAA formatting...")
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

    doc = Document(DOCX_OUT)

    # ── Global: ensure Times New Roman 12pt throughout ──────────────────────
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = "Times New Roman"
            if run.font.size is None or run.font.size < Pt(10):
                run.font.size = Pt(12)

        # ── Detect and reformat special paragraph types ──────────────────
        txt = para.text.strip()

        # Title paragraph
        if para.style.name in ("Title",) or (
            para.style.name.startswith("Heading") and
            "periodic structure" in txt.lower()
        ):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(14)
            continue

        # Abstract label
        if txt.lower().startswith("abstract"):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.left_indent  = Cm(1.0)
            para.paragraph_format.right_indent = Cm(1.0)
            for run in para.runs:
                run.font.italic = True
            continue

        # Keywords line
        if txt.lower().startswith("keywords"):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.bold = False
            # Bold just the "Keywords:" label
            if para.runs:
                para.runs[0].font.bold = True
            continue

        # Section headings (Heading 1)
        if para.style.name == "Heading 1":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after  = Pt(6)
            for run in para.runs:
                run.font.bold   = True
                run.font.italic = False
                run.font.size   = Pt(12)
                run.font.name   = "Times New Roman"
            continue

        # Subsection headings (Heading 2)
        if para.style.name == "Heading 2":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after  = Pt(3)
            for run in para.runs:
                run.font.bold   = True
                run.font.italic = True
                run.font.size   = Pt(12)
                run.font.name   = "Times New Roman"
            continue

        # Heading 3
        if para.style.name == "Heading 3":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.bold   = False
                run.font.italic = True
                run.font.size   = Pt(12)
                run.font.name   = "Times New Roman"
            continue

        # Body text: justified, single-spaced
        if para.style.name in ("Normal", "Body Text", "First Paragraph"):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para.paragraph_format.space_after  = Pt(6)
            para.paragraph_format.space_before = Pt(0)
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

    # ── Page setup ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.page_width    = Cm(21.0)
        section.page_height   = Cm(29.7)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    doc.save(DOCX_OUT)
    print(f"  Post-processing complete: {DOCX_OUT}")


# ── Step 5: Print a compliance summary ────────────────────────────────────────
def print_summary():
    from docx import Document
    doc = Document(DOCX_OUT)
    n_para  = len(doc.paragraphs)
    n_tables = len(doc.tables)
    # Rough word count
    word_count = sum(len(p.text.split()) for p in doc.paragraphs)

    print("\n" + "="*60)
    print("  MAA SUBMISSION DOCUMENT SUMMARY")
    print("="*60)
    print(f"  Output file  : {DOCX_OUT.name}")
    print(f"  Paragraphs   : {n_para}")
    print(f"  Tables       : {n_tables}")
    print(f"  ~Word count  : {word_count:,}")
    print()
    print("  MAA CHECKLIST (verify manually):")
    print("  [ ] Font: Times New Roman 12pt throughout")
    print("  [ ] Margins: 2.5 cm all sides, A4")
    print("  [ ] Abstract ≤ 250 words")
    print("  [ ] Keywords ≤ 8, no title words")
    print("  [ ] Section order: Intro → Methods → Results → Discussion → Conclusion")
    print("  [ ] References: alphabetical, MAA author-year format")
    print("  [ ] Figures & tables embedded inline (not at end)")
    print("  [ ] All \\macro placeholders resolved (check for backslash-words)")
    print("  [ ] Author Declaration Form attached separately")
    print("  [ ] Zenodo DOI verified live")
    print("="*60)
    print(f"\n  NEXT STEP: Open {DOCX_OUT.name} in Word/LibreOffice,")
    print("  compare side-by-side with the MAA template:")
    print("  https://docs.google.com/document/d/13ZVVMYa-YDAeUt6uf0JCPTNdmkPHrJ5h/edit")
    print("  Then email to: contact@maajournal.com\n")


# ── Main ───────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    os.chdir(MANUSCRIPT_DIR)
    compile_latex()
    build_reference_docx()
    run_pandoc()
    post_process_docx()
    print_summary()
